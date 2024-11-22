import os
from docx import Document
import PyPDF2


def read_file(file_path, include_contents=False):
    """
    Reads the contents of a file or the structure of a folder from the given path.
    Supports .docx (Word), .py (Python), and .pdf (PDF) files for file reading.

    If the path is a folder, it lists all files and optionally reads their contents.
    
    Parameters:
        file_path (str): Path to the file or folder.
        include_contents (bool): Whether to include the contents of files when reading a folder.
    
    Returns:
        dict or str: For folders, a dictionary with file paths as keys and contents (or None) as values.
                     For files, the file content as a string.
    """
    try:
        # Check if the path is a folder
        if os.path.isdir(file_path):
            folder_structure = {}
            for root, _, files in os.walk(file_path):
                for file in files:
                    abs_path = os.path.join(root, file)
                    if include_contents:
                        folder_structure[abs_path] = read_file(abs_path)  # Recursively read file contents
                    else:
                        folder_structure[abs_path] = None  # Only list the file path
            return folder_structure
        
        # If the path is a file, read its content
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".pdf"):
            pdf_text = ""
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    pdf_text += page.extract_text()
            return pdf_text.strip()
        elif file_path.endswith(".py"):
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        else:
            return f"Unsupported file type for reading: {file_path}"
    
    except FileNotFoundError:
        return f"Error: File not found at {file_path}"
    except Exception as e:
        return f"Error reading file: {e}"


def write_file(file_path, content):
    """
    Writes or updates a file with the specified content.
    Supports .docx (Word) and .py (Python) files.
    """
    try:
        if file_path.endswith(".docx"):
            # Write to Word (.docx) file
            doc = Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(file_path)
            return f"Word file successfully written to {file_path}"
        
        elif file_path.endswith(".py"):
            # Write to Python (.py) file as plain text
            with open(file_path, "w", encoding="utf-8") as file:
                file.write(content)
            return f"Python file successfully written to {file_path}"
        
        elif file_path.endswith(".pdf"):
            # PDF writing is not natively supported with PyPDF2, return an error
            return "Error: Writing to PDF files is not supported."
        
        else:
            return f"Unsupported file type for writing: {file_path}"
    
    except Exception as e:
        return f"Error writing file: {e}"