import os
import json
import logging
from docx import Document
from PyPDF2 import PdfReader
from LLM_interface.query_llm import (
    preprocess_prompt_with_functions,
    query_llm_function_decision,
    query_llm_marked_response
)
from Functions.local_formatter import LocalFormatter

def llm_decision(user_prompt: str, api_url: str = None, model_name: str = None):
    logging.info("llm_decision: Starting decision-making process.")
    enriched_prompt = preprocess_prompt_with_functions(user_prompt)
    logging.debug(f"llm_decision: Enriched prompt: {enriched_prompt}")

    llm_response = query_llm_function_decision(api_url, model_name, enriched_prompt, stream=False)
    logging.debug(f"llm_decision: Raw LLM Decision Response: {llm_response}")

    try:
        response_data = json.loads(llm_response)
        logging.info("llm_decision: Successfully parsed LLM response as JSON.")

        functions = response_data.get("function", [])
        parameters = response_data.get("parameters", [])

        if isinstance(functions, str):
            functions = [functions]
        if isinstance(parameters, dict):
            parameters = [parameters]
        if len(functions) != len(parameters):
            raise ValueError("Mismatch between number of functions and parameters.")

        results = []
        for func, param in zip(functions, parameters):
            logging.info(f"llm_decision: Processing function: {func} with parameters: {param}")
            path = param.get("path", "")

            if func == "handle_path":
                action_response = handle_path(path, api_url, model_name)
                for action_func, action_param in zip(action_response.get("function", []), action_response.get("parameters", [])):
                    if action_func == "read_file":
                        results.append(read_file(**action_param))
                    elif action_func == "list_folder":
                        results.append(list_folder(**action_param))
                    else:
                        results.append({"html_response": f"<p>Unknown action '{action_func}'</p>"})
            elif func == "read_file":
                results.append(read_file(**param))
            elif func == "write_file":
                results.append(write_file(**param))
            elif func == "list_folder":
                results.append(list_folder(**param))
            elif func == "general_question":
                # This should return a generator for streaming
                gen = general_question(param.get("general_question", ""), api_url, model_name, stream=True)
                results.append({"stream_generator": gen})
            else:
                logging.warning(f"llm_decision: Unknown function '{func}'.")
                results.append({"html_response": f"<p>Error: Unknown function '{func}'</p>"})

        combined_html_response = ""
        stream_generator = None
        combined_detailed_info = []
        for res in results:
            if "stream_generator" in res:
                logging.debug("llm_decision: Found stream_generator in results.")
                stream_generator = res["stream_generator"]
            else:
                html_response = res.get("html_response")
                if html_response:
                    combined_html_response += html_response
                detailed_info = res.get("detailed_info", {})
                if detailed_info:
                    combined_detailed_info.append(detailed_info)

        if stream_generator:
            logging.info("llm_decision: Returning stream generator for streaming response.")
            return {
                "stream_generator": stream_generator,
                "detailed_info": combined_detailed_info
            }
        else:
            logging.info("llm_decision: Returning standard HTML response.")
            return {
                "html_response": combined_html_response,
                "detailed_info": combined_detailed_info,
            }

    except json.JSONDecodeError:
        logging.error(f"llm_decision: Error decoding LLM response as JSON: {llm_response}")
        return {
            "html_response": f"<p>Error: Invalid JSON response from LLM. Raw output: {llm_response}</p>",
            "detailed_info": {},
        }
    except Exception as e:
        logging.error(f"llm_decision: Error in llm_decision function: {e}")
        return {
            "html_response": f"<p>Error executing function: {e}</p>",
            "detailed_info": {},
        }

def handle_path(path, api_url, model_name):
    logging.info(f"handle_path: Handling path: {path}")
    if os.path.isfile(path):
        return {
            "function": ["read_file"],
            "parameters": [{"path": path}]
        }
    elif os.path.isdir(path):
        return {
            "function": ["list_folder"],
            "parameters": [{"path": path}]
        }
    else:
        return {
            "function": ["error"],
            "parameters": [
                {"error_message": f"{path} is neither a valid file nor a folder."}
            ]
        }

def read_file(path: str):
    logging.info(f"read_file: Reading file at path: {path}")
    try:
        file_metadata = {"name": os.path.basename(path), "contents": None}
        if path.endswith(".docx"):
            doc = Document(path)
            file_metadata["contents"] = "\n".join([para.text for para in doc.paragraphs])
        elif path.endswith(".pdf"):
            pdf_text = ""
            with open(path, "rb") as pdf_file:
                pdf_reader = PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    pdf_text += page.extract_text() or ""
            file_metadata["contents"] = pdf_text.strip()
        elif any(path.endswith(ext) for ext in [".py", ".md", ".txt"]):
            with open(path, "r", encoding="utf-8") as f:
                file_metadata["contents"] = f.read()
        else:
            return {
                "plain_text_response": f"Unsupported file type: {file_metadata['name']}",
                "detailed_info": file_metadata
            }

        if not file_metadata["contents"]:
            raise Exception("Failed to extract file content.")

        enriched_prompt = f"Explain the following content:\n\n{file_metadata['contents']}"
        logging.debug("read_file: Enriched prompt prepared. The explanation will be handled separately if needed.")
        # Since read_file does not know api_url/model_name from this scope directly,
        # we'll handle explanation separately if needed by the LLM decision logic.
        return {
            "plain_text_response": "File content read successfully. Use 'general_question' function to explain.",
            "detailed_info": file_metadata
        }

    except FileNotFoundError:
        logging.error(f"read_file: File not found at {path}")
        return {
            "html_response": f"<p>Error: File not found at {path}</p>",
            "detailed_info": {"name": os.path.basename(path)}
        }
    except Exception as e:
        logging.error(f"read_file: Error reading file: {e}")
        return {
            "html_response": f"<p>Error reading file: {e}</p>",
            "detailed_info": {"name": os.path.basename(path)}
        }

def write_file(path: str, content: str):
    logging.info(f"write_file: Writing to file at path: {path}")
    try:
        if path.endswith(".docx"):
            doc = Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(path)
        elif any(path.endswith(ext) for ext in [".py", ".md", ".txt"]):
            with open(path, "w", encoding="utf-8") as file:
                file.write(content)
        elif path.endswith(".pdf"):
            return {
                "plain_text": "Error: Writing to PDF files is not supported.",
                "detailed_info": {"path": path, "file_type": "pdf"},
            }
        else:
            return {
                "plain_text": f"Unsupported file type for writing: {path}",
                "detailed_info": {"path": path},
            }

        logging.info("write_file: Successfully wrote to file.")
        return {"plain_text": f"File successfully written to {path}", "detailed_info": {"path": path}}
    except Exception as e:
        logging.error(f"write_file: Error writing file: {e}")
        return {"plain_text": f"Error writing file: {e}", "detailed_info": {"path": path}}

def list_folder(path: str) -> dict:
    logging.info(f"list_folder: Listing folder at path: {path}")
    import os
    from html import escape

    def read_file_content(file_path: str) -> str:
        try:
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            if ext == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
            elif ext in {'.txt', '.py', '.js', '.html', '.md'}:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif ext == '.docx':
                d = Document(file_path)
                return "\n".join(paragraph.text for paragraph in d.paragraphs)
            return "Unsupported file type for preview."
        except Exception as e:
            logging.error(f"list_folder: Error reading file {file_path}: {e}")
            return f"Error reading file: {e}"

    def build_tree_and_content(path, level=0):
        indent = "    " * level
        tree_prefix = f"{indent}├── "
        html_indent = "&nbsp;&nbsp;&nbsp;&nbsp;" * level
        plain_text_structure = ""
        html_tree_structure = ""
        aggregated_content = ""

        if os.path.isdir(path):
            dir_name = os.path.basename(path)
            plain_text_structure += f"{tree_prefix}{dir_name}/\n"
            html_tree_structure += f"{html_indent}<strong>{escape(dir_name)}/</strong><br>"
            for item in sorted(os.listdir(path)):
                if item.startswith('.') or item == "__pycache__":
                    continue
                item_path = os.path.join(path, item)
                sub_plain, sub_html, sub_content = build_tree_and_content(item_path, level + 1)
                plain_text_structure += sub_plain
                html_tree_structure += sub_html
                aggregated_content += sub_content
        else:
            file_name = os.path.basename(path)
            plain_text_structure += f"{tree_prefix}{file_name}\n"
            html_tree_structure += f"{html_indent}{escape(file_name)}<br>"
            file_content = read_file_content(path)
            aggregated_content += f"\n# {file_name}\n{file_content}\n"

        return plain_text_structure, html_tree_structure, aggregated_content

    # Path validation
    if not os.path.exists(path):
        logging.error("list_folder: Path does not exist.")
        return {"html_response": "<p>Error: Path does not exist.</p>", "detailed_info": {"error": "Path does not exist"}}
    if not os.path.isdir(path):
        logging.error("list_folder: Path is not a folder.")
        return {"html_response": "<p>Error: Path is not a folder.</p>", "detailed_info": {"error": "Path is not a folder"}}

    try:
        folder_structure, html_tree_structure, aggregated_content = build_tree_and_content(path)

        # Prepare the prompt for explanation
        enriched_prompt = (
            f"Here is the folder structure of a programming project:\n\n{folder_structure}\n\n"
            f"Below are the contents of the files:\n{aggregated_content}\n\n"
            "Explain the overall purpose of the project, key components, and functionality."
        )

        # We'll let llm_decision handle calling general_question as needed
        return {
            "plain_text_response": "Folder structure read successfully. Use 'general_question' function to explain.",
            "detailed_info": {"folder_structure": folder_structure, "explanation_prompt": enriched_prompt}
        }

    except Exception as e:
        logging.error(f"list_folder: Error processing folder: {e}")
        return {"html_response": f"<p>Error: {e}</p>", "detailed_info": {"error": str(e)}}

def general_question(user_prompt, api_url, model_name, stream=False):
    logging.info(f"general_question: Handling prompt: {user_prompt}, stream={stream}")
    response_chunks = query_llm_marked_response(api_url, model_name, user_prompt, stream=stream)

    if stream:
        formatter = LocalFormatter()
        chunk_index = 0
        try:
            for chunk in response_chunks:
                chunk_index += 1
                logging.debug(f"general_question: Received chunk #{chunk_index}: {chunk[:100]}...")
                html = formatter.feed_text(chunk)
                if html:
                    logging.debug(f"general_question: Yielding formatted HTML chunk of length {len(html)}.")
                    yield html
            final_html = formatter.close()
            if final_html:
                logging.debug("general_question: Yielding final formatted HTML after close.")
                yield final_html
        except Exception as e:
            logging.error(f"general_question: Error during streaming: {e}")
            yield f"<p>Error during streaming: {e}</p>"
    else:
        all_text = ""
        for chunk in response_chunks:
            all_text += chunk
        formatter = LocalFormatter()
        html = formatter.feed_text(all_text)
        html += formatter.close()
        return {
            "html_response": html,
            "detailed_info": {"type": "general_question", "prompt": user_prompt},
        }